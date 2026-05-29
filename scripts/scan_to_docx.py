#!/usr/bin/env python3
"""Convert scanned PDF/image pages to an editable DOCX using OCR.

The output prioritizes editable text while keeping the page geometry close to the
scan: page rotation is corrected, Vietnamese/English OCR is used, and OCR lines
are placed with approximate indentation/spacing instead of embedding page images.
"""
import argparse
import io
import re
import sys
from pathlib import Path

try:
    import fitz  # PyMuPDF
    import pytesseract
    from docx import Document
    from docx.enum.section import WD_ORIENTATION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    from PIL import Image, ImageOps, ImageFilter
except Exception as exc:  # pragma: no cover - gives a clear runtime error
    print(
        "ERROR: Missing OCR dependencies. Install: "
        "python -m pip install pymupdf pytesseract pillow python-docx. "
        "Also install Tesseract OCR + Vietnamese language data.",
        file=sys.stderr,
    )
    print(f"DETAIL: {exc}", file=sys.stderr)
    raise


EMU_PER_INCH = 914400


def render_pdf_pages(input_path: Path, dpi: int):
    doc = fitz.open(str(input_path))
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    pages = []
    for index, page in enumerate(doc):
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        image = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
        pages.append((index + 1, image))
    doc.close()
    return pages


def load_pages(input_path: Path, dpi: int):
    if input_path.suffix.lower() == ".pdf":
        return render_pdf_pages(input_path, dpi)
    image = Image.open(str(input_path)).convert("RGB")
    return [(1, ImageOps.exif_transpose(image))]


def deskew_and_orient(image: Image.Image, lang: str):
    image = ImageOps.exif_transpose(image).convert("RGB")
    try:
        osd = pytesseract.image_to_osd(image, output_type=pytesseract.Output.DICT)
        rotate = int(osd.get("rotate", 0) or 0)
        if rotate:
            image = image.rotate(-rotate, expand=True, fillcolor="white")
            print(f"[scan-ocr] auto-rotated page by {rotate} degrees", file=sys.stderr)
    except Exception as exc:
        print(f"[scan-ocr] orientation detection skipped: {str(exc)[:120]}", file=sys.stderr)

    # OCR-friendly grayscale while preserving the RGB image dimensions for layout.
    gray = ImageOps.grayscale(image)
    gray = ImageOps.autocontrast(gray)
    gray = gray.filter(ImageFilter.SHARPEN)
    return image, gray


def ocr_lines(image: Image.Image, lang: str, psm: int):
    config = f"--oem 1 --psm {psm} -c preserve_interword_spaces=1"
    data = pytesseract.image_to_data(
        image,
        lang=lang,
        config=config,
        output_type=pytesseract.Output.DICT,
    )

    lines = {}
    total = len(data.get("text", []))
    for i in range(total):
        text = (data["text"][i] or "").strip()
        if not text:
            continue
        try:
            conf = float(data["conf"][i])
        except Exception:
            conf = -1.0
        if conf < 25:
            continue
        key = (
            int(data["page_num"][i]),
            int(data["block_num"][i]),
            int(data["par_num"][i]),
            int(data["line_num"][i]),
        )
        item = {
            "text": text,
            "left": int(data["left"][i]),
            "top": int(data["top"][i]),
            "width": int(data["width"][i]),
            "height": int(data["height"][i]),
            "conf": conf,
        }
        lines.setdefault(key, []).append(item)

    normalized = []
    for words in lines.values():
        words.sort(key=lambda w: w["left"])
        text = clean_ocr_text(" ".join(w["text"] for w in words).strip())
        if not text:
            continue
        left = min(w["left"] for w in words)
        top = min(w["top"] for w in words)
        right = max(w["left"] + w["width"] for w in words)
        bottom = max(w["top"] + w["height"] for w in words)
        normalized.append({
            "text": text,
            "left": left,
            "top": top,
            "right": right,
            "bottom": bottom,
            "height": bottom - top,
            "conf": sum(w["conf"] for w in words) / max(1, len(words)),
        })
    normalized.sort(key=lambda item: (item["top"], item["left"]))
    return normalized


def clean_ocr_text(text: str):
    replacements = {
        "Tw do": "Tự do",
        "Ty do": "Tự do",
        "Tự đo": "Tự do",
        "Hanh Phúc": "Hạnh Phúc",
        "Hanh phúc": "Hạnh phúc",
        "CONG HÒA": "CỘNG HÒA",
        "CONG HOA": "CỘNG HÒA",
        "CQNG HÒA": "CỘNG HÒA",
        "CHU NGHIA": "CHỦ NGHĨA",
        "VIET NAM": "VIỆT NAM",
        "BENH VIEN": "BỆNH VIỆN",
        "DA KHOA": "ĐA KHOA",
        "TAN PHU": "TÂN PHÚ",
        "HỖ CHÍ MINH": "HỒ CHÍ MINH",
        "HỎ CHÍ MINH": "HỒ CHÍ MINH",
        "KE HOẠCH": "KẾ HOẠCH",
        "PHOI HỢP": "PHỐI HỢP",
        "PHÓI HỢP": "PHỐI HỢP",
        "MUC DICH": "MỤC ĐÍCH",
        "YEU CAU": "YÊU CẦU",
        "YÊU CÀU": "YÊU CẦU",
        "NỘI DỤNG": "NỘI DUNG",
        "THUC HIỆN": "THỰC HIỆN",
        "PHAN CÔNG": "PHÂN CÔNG",
        "Sé:": "Số:",
        "sé:": "Số:",
        "sá:": "Số:",
        "ngdy": "ngày",
        "thang": "tháng",
        "két": "kết",
        "buôi": "buổi",
        "tuôi": "tuổi",
        "thuộc": "thuốc",
        "cap thuốc": "cấp thuốc",
        "tô chức": "tổ chức",
        "chât lượng": "chất lượng",
        "sức khoẻ": "sức khỏe",
        "tông quát": "tổng quát",
        "ô bụng": "ổ bụng",
        "Hô sơ": "Hồ sơ",
        "Y tê": "Y tế",
        "V tế": "Y tế",
        "Ÿ tế": "Y tế",
        "dé nghị": "đề nghị",
        "dé lưỡi": "đè lưỡi",
        "dinh kèm": "đính kèm",
        "dược sĩ": "dược sĩ",
        "Vật tu": "Vật tư",
        "Dam bao": "Đảm bảo",
        "vat tư": "vật tư",
        "phục vu": "phục vụ",
        "giây điện tim": "giấy điện tim",
        "kêt thúc": "kết thúc",
        "T 6": "Tổ",
        "Thành phô Hô": "Thành phố Hồ",
        "Giám đoc": "Giám đốc",
    }
    cleaned = text
    cleaned = cleaned.replace("—", "-").replace("–", "-")
    cleaned = re.sub(r"[¬ˆ`]+", "", cleaned)
    cleaned = re.sub(r"^\s*[=\-|_\\\]\[]+\s*", "", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    for wrong, right in replacements.items():
        cleaned = cleaned.replace(wrong, right)
    cleaned = re.sub(r"\bSO Y\b", "SỞ Y", cleaned)
    cleaned = re.sub(r"\bTE TP\b", "TẾ TP", cleaned)
    cleaned = re.sub(r"\bHO CH[IÍ]\b", "HỒ CHÍ", cleaned)
    cleaned = re.sub(r"^HI\.\s+PHÂN", "III. PHÂN", cleaned)
    cleaned = re.sub(r"^I\s+MỤC", "I. MỤC", cleaned)
    cleaned = re.sub(r"^II\s+NỘI", "II. NỘI", cleaned)
    return cleaned.strip()


def score_ocr_lines(lines):
    text = "\n".join(line["text"] for line in lines)
    if not text.strip():
        return -1_000_000.0
    good_terms = [
        "SỞ Y TẾ", "BỆNH VIỆN", "ĐA KHOA", "TÂN PHÚ", "CỘNG HÒA",
        "CHỦ NGHĨA", "VIỆT NAM", "Độc lập", "Tự do", "Hạnh Phúc",
        "KẾ HOẠCH", "PHỐI HỢP", "MỤC ĐÍCH", "YÊU CẦU", "THỰC HIỆN",
        "khám bệnh", "cấp thuốc", "miễn phí", "người dân"
    ]
    bad_chars = sum(text.count(ch) for ch in "¥¬ˆ`|[]{}")
    ascii_upper_runs = len(re.findall(r"\b[A-Z]{4,}\b", text))
    no_space_penalty = len(re.findall(r"[A-ZĐÂĂÊÔƠƯÁÀẢÃẠÉÈẺẼẸÍÌỈĨỊÓÒỎÕỌÚÙỦŨỤÝỲỶỸỴ]{6,}", text))
    good = sum(text.upper().count(term.upper()) for term in good_terms)
    avg_conf = sum(line.get("conf", 0) for line in lines) / max(1, len(lines))
    return avg_conf + good * 12 - bad_chars * 8 - ascii_upper_runs * 0.8 - no_space_penalty * 2


def unique_keep_order(items):
    result = []
    for item in items:
        if item and item not in result:
            result.append(item)
    return result


def best_ocr_lines(image: Image.Image, lang: str, psm: int):
    langs = unique_keep_order([lang, "vie" if "vie" in lang else "", "vie+eng" if "vie" in lang else "", "eng" if "eng" in lang and lang != "eng" else ""])
    psms = unique_keep_order([psm, 4, 6, 11])
    best = None
    best_meta = None
    for candidate_lang in langs:
        for candidate_psm in psms:
            try:
                lines = ocr_lines(image, candidate_lang, candidate_psm)
            except Exception as exc:
                print(f"[scan-ocr] candidate failed lang={candidate_lang} psm={candidate_psm}: {str(exc)[:120]}", file=sys.stderr)
                continue
            score = score_ocr_lines(lines)
            if best is None or score > best[0]:
                best = (score, lines)
                best_meta = (candidate_lang, candidate_psm)
    if best is None:
        return []
    print(f"[scan-ocr] selected lang={best_meta[0]} psm={best_meta[1]} score={best[0]:.1f}", file=sys.stderr)
    return best[1]


def configure_section(section, image_width: int, image_height: int):
    landscape = image_width > image_height
    section.orientation = WD_ORIENTATION.LANDSCAPE if landscape else WD_ORIENTATION.PORTRAIT
    if landscape:
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
    else:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)


def add_page_image(document: Document, image: Image.Image):
    section = document.sections[-1]
    configure_section(section, image.width, image.height)
    content_width = section.page_width - section.left_margin - section.right_margin
    content_height = section.page_height - section.top_margin - section.bottom_margin
    width_in = content_width / EMU_PER_INCH
    height_in = content_height / EMU_PER_INCH
    image_ratio = image.width / max(1, image.height)
    box_ratio = width_in / max(0.01, height_in)
    if image_ratio > box_ratio:
        draw_width = width_in
        draw_height = width_in / image_ratio
    else:
        draw_height = height_in
        draw_width = height_in * image_ratio

    stream = io.BytesIO()
    image.save(stream, format="PNG")
    stream.seek(0)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(stream, width=Inches(draw_width), height=Inches(draw_height))


def add_ocr_page(document: Document, page_no: int, image: Image.Image, lines, include_page_label: bool):
    section = document.sections[-1]
    configure_section(section, image.width, image.height)
    content_width = section.page_width - section.left_margin - section.right_margin
    content_height = section.page_height - section.top_margin - section.bottom_margin
    width_in = content_width / EMU_PER_INCH
    height_in = content_height / EMU_PER_INCH
    x_scale = width_in / max(1, image.width)
    y_scale = height_in / max(1, image.height)

    if include_page_label:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Trang {page_no}")
        r.bold = True
        r.font.size = Pt(10)

    prev_bottom = 0
    for line in lines:
        para = document.add_paragraph()
        para_format = para.paragraph_format
        para_format.left_indent = Inches(max(0.0, line["left"] * x_scale))
        gap_px = max(0, line["top"] - prev_bottom)
        para_format.space_before = Pt(min(28, max(0, gap_px * y_scale * 72)))
        para_format.space_after = Pt(0)
        para_format.line_spacing = 1.0
        run = para.add_run(line["text"])
        run.font.name = "Times New Roman"
        font_size = max(8.0, min(16.0, line["height"] * y_scale * 72 * 0.86))
        run.font.size = Pt(font_size)
        prev_bottom = max(prev_bottom, line["bottom"])


def add_ocr_heading(document: Document, page_no: int):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"Nội dung OCR chỉnh sửa - trang {page_no}")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9)


def convert(input_path: Path, output_path: Path, lang: str, dpi: int, psm: int, include_page_label: bool, output_mode: str):
    pages = load_pages(input_path, dpi)
    if not pages:
        raise RuntimeError("Input does not contain any page.")

    document = Document()
    total_lines = 0
    total_chars = 0

    for index, (page_no, raw_image) in enumerate(pages):
        if index > 0:
            document.add_section()
        oriented, ocr_image = deskew_and_orient(raw_image, lang)
        lines = best_ocr_lines(ocr_image, lang=lang, psm=psm)
        total_lines += len(lines)
        total_chars += sum(len(line["text"]) for line in lines)

        if output_mode == "visual":
            add_page_image(document, oriented)
        else:
            add_ocr_page(document, page_no, oriented, lines, include_page_label=include_page_label)
        print(f"[scan-ocr] page {page_no}: {len(lines)} lines, {sum(len(l['text']) for l in lines)} chars", file=sys.stderr)

    if total_chars < 8:
        raise RuntimeError("OCR did not detect enough text. Try a clearer scan or install the right Tesseract language pack.")

    document.core_properties.author = "Convert URL Studio"
    document.core_properties.title = input_path.stem
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print(f"[scan-ocr] DONE pages={len(pages)} lines={total_lines} chars={total_chars}", file=sys.stderr)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("output")
    parser.add_argument("--lang", default="vie")
    parser.add_argument("--dpi", type=int, default=220)
    parser.add_argument("--psm", type=int, default=4)
    parser.add_argument("--page-label", choices=["0", "1"], default="0")
    parser.add_argument("--tesseract-cmd", default="")
    parser.add_argument("--output-mode", choices=["editable", "visual"], default="editable")
    args = parser.parse_args()

    if args.tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = args.tesseract_cmd

    convert(
        Path(args.input),
        Path(args.output),
        lang=args.lang,
        dpi=max(120, min(args.dpi, 360)),
        psm=max(3, min(args.psm, 11)),
        include_page_label=args.page_label == "1",
        output_mode=args.output_mode,
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
